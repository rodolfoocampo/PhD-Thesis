from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

# Title
title = doc.add_heading('Citation Cluster Changes — Before & After', level=1)

clusters = [
    {
        "num": 3,
        "before": r"Since then, across a growing literature, authors argue for the potential of human-AI co-creativity as a scenario that can augment, enhance and support human creativity, mitigating many of the risks posed by this technology while unlocking the beneficial opportunities it provides. \citep{Yannakakis2014-zs,Kantosalo2020-zf,Rezwana2022-gg,Moruzzi2024-cq,Haase2024-yp,Lin2023-zq,Karimi2018-wi,Vinchon2023-gh}.",
        "after": r"Since then, across a growing literature, authors argue for the potential of human-AI co-creativity as a scenario that can augment, enhance and support human creativity, mitigating many of the risks posed by this technology while unlocking the beneficial opportunities it provides. Foundational work has defined the concept and explored synergies across different levels of collaboration \citep{Yannakakis2014-zs, Haase2024-yp, Vinchon2023-gh}, while several frameworks and ontologies have been proposed to classify and describe co-creative systems \citep{Kantosalo2020-zf, Rezwana2022-gg, Lin2023-zq, Moruzzi2024-cq}, alongside methods for evaluating creativity within them \citep{Karimi2018-wi}.",
    },
    {
        "num": 4,
        "before": r"Since then, many researchers have sought to describe how computers could act creatively and how this creativity can be measured \citep{Boden2003-hk, Boden1998-yn, Colton2012-jc, Bown2012-gg, Moruzzi2020-mw, Wiggins2006-zd, Jordanous2012-kw}.",
        "after": r"Since then, many researchers have sought to describe how computers could act creatively and how this creativity can be measured. Foundational theoretical work explored the nature of artificial creativity \citep{Boden1998-yn, Boden2003-hk}, leading to formal frameworks for describing and evaluating creative systems \citep{Wiggins2006-zd, Jordanous2012-kw}. Others have investigated creative AI systems and the philosophical implications of machine creativity \citep{Colton2012-jc, Bown2012-gg, Moruzzi2020-mw}.",
    },
    {
        "num": 5,
        "before": r"These conflicting findings suggest that outcomes depend heavily on how users interact with AI, which can be shaped by interface and interaction design but also personal motivations and context \citep{Lehmann2022-kr, Kantosalo2020-nh,Liapis2016-bv,Lin2023-jd,Karimi2020-cf,Moruzzi2024-cq, Rezwana2022-gg,Abbas2024-sf}.",
        "after": r"These conflicting findings suggest that outcomes depend heavily on how users interact with AI, which can be shaped by interface and interaction design but also personal motivations and context. Interaction frameworks for co-creativity have explored how different design choices affect collaboration dynamics \citep{Kantosalo2020-nh, Rezwana2022-gg, Moruzzi2024-cq}, while empirical work has examined the design space and specific modalities that shape user engagement \citep{Lehmann2022-kr, Liapis2016-bv, Lin2023-jd, Karimi2020-cf, Abbas2024-sf}.",
    },
    {
        "num": 6,
        "before": r"Since then, a growing literature frames human-AI co-creativity as a type of interaction that ultimately augments human creativity and can mitigate risks associated with the automation of creative work \citep{Yannakakis2014-zs, Kantosalo2020-zf, Rezwana2022-gg, Moruzzi2024-cq, Haase2024-yp, Lin2023-zq, Karimi2018-wi}.",
        "after": r"Since then, a growing literature frames human-AI co-creativity as a type of interaction that ultimately augments human creativity and can mitigate risks associated with the automation of creative work. Early work established the conceptual foundations of mixed-initiative co-creativity \citep{Yannakakis2014-zs, Haase2024-yp}, while subsequent frameworks have classified interaction dynamics and design dimensions for co-creative systems \citep{Kantosalo2020-zf, Rezwana2022-gg, Lin2023-zq, Moruzzi2024-cq}, alongside methods for evaluating the creativity of such systems \citep{Karimi2018-wi}.",
    },
    {
        "num": 7,
        "before": r"...how creative behaviour of computers can be measured \citep{Ritchie2007-jy, Colton2008-fh, Colton2011-uy, Maher2012-oj, Jordanous2012-kw}, as well as rich creative practices that employ computers both as tools, and as autonomous and collaborative creative agents \citep{Cohen1995-wt, Colton2015-qr, Perez-y-Perez1999-ma, Cope1992-pq, Reichardt1968-eo}.",
        "after": r"...and how creative behaviour of computers can be measured. Theoretical work has proposed formal criteria for attributing creativity to programs \citep{Ritchie2007-jy, Jordanous2012-kw}, descriptive models such as FACE and IDEA \citep{Colton2008-fh, Colton2011-uy}, and examinations of creative attribution in computational and collective processes \citep{Maher2012-oj}. Alongside this theoretical work, rich creative practices have employed computers as creative agents, from early computer art \citep{Reichardt1968-eo} and autonomous painting systems \citep{Cohen1995-wt, Colton2015-qr} to computational story generation \citep{Perez-y-Perez1999-ma} and music composition \citep{Cope1992-pq}.",
    },
    {
        "num": 8,
        "before": r"Numerous fields within HCI have established principles for designing effective human-computer interaction and tools \citep{Nielsen1994-df, Amershi2019-vy, Shneiderman2020-je, Wright2020-nt, Bengler2012-jf, Resnick2005-fs}.",
        "after": r"Numerous fields within HCI have established principles for designing effective human-computer interaction and tools, from foundational usability heuristics \citep{Nielsen1994-df} to guidelines specifically addressing human-AI interaction \citep{Amershi2019-vy, Shneiderman2020-je, Wright2020-nt}, as well as principles for cooperative systems \citep{Bengler2012-jf} and creativity support tools \citep{Resnick2005-fs}.",
    },
    {
        "num": 9,
        "before": r"Second, observed metrics that assess task completion time, quality, creativity, and quantity of outputs, often rated by crowdsourced evaluators or experts \citep{Kim2021-fh, Kantosalo2019-pz, Rezwana2022-ui, Rezwana2023-gj, Lawton2023-tb}.",
        "after": r"Second, observed metrics that assess task completion time, quality, creativity, and quantity of outputs, often rated by crowdsourced evaluators or experts. These metrics have been employed across studies evaluating co-creative design and ideation \citep{Kim2021-fh, Kantosalo2019-pz}, user perceptions and interaction design comparisons \citep{Rezwana2022-ui, Lawton2023-tb}, and ethical dimensions of co-creativity \citep{Rezwana2023-gj}.",
    },
    {
        "num": 10,
        "before": r"...with users commonly reporting concerns with the lack of originality, generic outputs, and cliched styles \citep{Chakrabarty2024-ov, Chang2023-tv, Clark1998-yi, Ippolito2022-mf, Li2024-yh}.",
        "after": r"...with users commonly reporting concerns with the lack of originality, generic outputs, and cliched styles. Professional writers have consistently flagged these limitations \citep{Chakrabarty2024-ov, Ippolito2022-mf}, as have creative practitioners developing expertise with prompt-based tools \citep{Chang2023-tv, Li2024-yh}.",
    },
    {
        "num": 11,
        "before": r"In contrast, there is growing interest in designing what can be understood as dialogic interaction sometimes referred explicitly as such, and sometimes describes merely as a back-and-forth process characterised by mutual influence, iterative refinement, and feedback loops \citep{Bown2020-oc, Gomez2023-bp, Wang2021-uy, Zhou2024-vp, Ghajargar2022-af, Feldman2017-ip}.",
        "after": r"In contrast, there is growing interest in designing what can be understood as dialogic interaction, sometimes referred to explicitly as such, and sometimes described merely as a back-and-forth process characterised by mutual influence, iterative refinement, and feedback loops. Some authors have explored dialogue and mutual understanding as foundations for co-creation \citep{Bown2020-oc, Wang2021-uy}, while others have proposed frameworks for capturing the nonlinear and iterative nature of human-AI collaboration \citep{Gomez2023-bp, Zhou2024-vp}. Experiential accounts of creative practice further illustrate how these reciprocal dynamics play out in writing and other domains \citep{Ghajargar2022-af, Feldman2017-ip}.",
    },
    {
        "num": 12,
        "before": r"Explainable AI is an active field focused on this issue \citep{Zhu2018-zd, Llano2022-ti, Newn2020-mv, Shneiderman2020-je, Linardatos2020-uq, El-Assady2022-qc, Gomez2023-bp}.",
        "after": r"Explainable AI is an active field focused on this issue. Within co-creative contexts specifically, frameworks such as XAID have addressed interpretability for designers \citep{Zhu2018-zd}, and arguments have been made for explainability as a prerequisite for deeper creative collaboration \citep{Llano2022-ti}. More broadly, comprehensive reviews of interpretability methods \citep{Linardatos2020-uq} and guidelines for trustworthy AI \citep{Shneiderman2020-je} provide a foundation, while research on communication channels \citep{Newn2020-mv, Gomez2023-bp} and cognitive biases triggered by explanations \citep{El-Assady2022-qc} highlights the complexity of designing effective transparency mechanisms.",
    },
    {
        "num": 13,
        "before": r"For instance, many studies find generative AI useful for early-stage ideation and brainstorming, helping users overcome blocks and explore possibilities \citep{Calderwood2020-gg, Clark2018-yf, Wan2023-he, Mirowski2023-oz, Doshi2023-dv}.",
        "after": r'For instance, many studies find generative AI useful for early-stage ideation and brainstorming, helping users overcome blocks and explore possibilities. Studies with writers have found that AI suggestions can spark new directions and serve as a creative catalyst \citep{Calderwood2020-gg, Clark2018-yf, Mirowski2023-oz}, while AI has also been described as a "second mind" during prewriting \citep{Wan2023-he}, though it may reduce the overall diversity of creative output \citep{Doshi2023-dv}.',
    },
    {
        "num": 14,
        "before": r"Providing a scaffolding throughout a creative process, particularly for novices has been increasingly explored in the literature, with promising outcomes \citep{Yuan2022-kb, Fan2019-qq, Ding2024-ta, Long2019-lw, Louie2020-aq, Ippolito2022-mf, Wadinambiarachchi2024-jn}.",
        "after": r"Providing scaffolding throughout a creative process, particularly for novices, has been increasingly explored in the literature, with promising outcomes. Writing tools have demonstrated how structured AI assistance can guide users through creative tasks \citep{Yuan2022-kb, Ippolito2022-mf}, while similar approaches have been applied to collaborative drawing \citep{Fan2019-qq} and novice music co-creation \citep{Louie2020-aq}. Research has also examined how GUI design can scaffold generative AI interactions at varying levels of task complexity \citep{Ding2024-ta}, how co-creative AI can be made accessible in public spaces \citep{Long2019-lw}, and the effects of AI-generated stimuli on divergent thinking during ideation \citep{Wadinambiarachchi2024-jn}.",
    },
]

RED = RGBColor(0xC0, 0x39, 0x2B)
GREEN = RGBColor(0x1E, 0x84, 0x49)
DARK = RGBColor(0x22, 0x22, 0x22)
GREY = RGBColor(0x66, 0x66, 0x66)

for i, c in enumerate(clusters):
    # Cluster heading
    h = doc.add_heading(f'Cluster {i+1}', level=2)

    # Before
    p = doc.add_paragraph()
    run = p.add_run('Before')
    run.bold = True
    run.font.color.rgb = RED
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run(c["before"])
    run.font.size = Pt(10)
    run.font.color.rgb = DARK

    # After
    p = doc.add_paragraph()
    run = p.add_run('After')
    run.bold = True
    run.font.color.rgb = GREEN
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run(c["after"])
    run.font.size = Pt(10)
    run.font.color.rgb = DARK

    # Separator
    doc.add_paragraph('—' * 40)

out = '/Users/rodolfoocampo/Desktop/31Mar/PhD-Thesis/citation-cluster-review.docx'
doc.save(out)
print(f'Saved to {out}')
